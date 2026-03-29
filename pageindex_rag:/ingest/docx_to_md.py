from docx import Document
from pathlib import Path


def _is_all_bold(paragraph) -> bool:
    """Return True if every non-empty run in the paragraph is bold."""
    runs_with_text = [r for r in paragraph.runs if r.text.strip()]
    if not runs_with_text:
        return False
    return all(r.bold for r in runs_with_text)


def _table_to_markdown(table) -> str:
    """Convert a docx table to a markdown table."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ""

    # Build markdown table
    lines = []
    header = rows[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    for row in rows[1:]:
        # Pad row if it has fewer cells than header
        padded = row + [""] * (len(header) - len(row))
        lines.append("| " + " | ".join(padded[:len(header)]) + " |")
    return "\n".join(lines)


def docx_to_markdown(path: str) -> str:
    doc = Document(path)
    lines = []

    # Build a map of element positions so we can interleave paragraphs and tables
    # in document order.
    from docx.oxml.ns import qn
    body = doc.element.body
    para_idx = 0
    table_idx = 0
    paragraphs = doc.paragraphs
    tables = doc.tables

    for child in body:
        if child.tag == qn("w:p"):
            if para_idx < len(paragraphs):
                p = paragraphs[para_idx]
                para_idx += 1
                text = p.text.strip()
                if not text:
                    continue

                style = p.style.name.lower()

                # Explicit Word heading styles
                if "heading 1" in style:
                    lines.append(f"# {text}")
                elif "heading 2" in style:
                    lines.append(f"## {text}")
                elif "heading 3" in style:
                    lines.append(f"### {text}")
                elif "heading 4" in style:
                    lines.append(f"#### {text}")
                # Custom heading styles (e.g. "MainHeading", "SubHeading")
                elif "mainheading" in style.replace(" ", ""):
                    lines.append(f"## {text}")
                elif "subheading" in style.replace(" ", ""):
                    lines.append(f"### {text}")
                # Bold-only short paragraphs that look like section titles
                elif _is_all_bold(p) and len(text) <= 120 and not text.endswith("."):
                    lines.append(f"## {text}")
                # List items
                elif style.startswith("list"):
                    lines.append(f"- {text}")
                else:
                    lines.append(text)
        elif child.tag == qn("w:tbl"):
            if table_idx < len(tables):
                table = tables[table_idx]
                table_idx += 1
                md_table = _table_to_markdown(table)
                if md_table:
                    lines.append(md_table)

    return "\n\n".join(lines)